"""生成虚构知识库素材：产品手册.docx + 售后政策.pdf（FAQ.html 为手工网页）。
虚构公司：云析科技 / 产品：智能会议平板 X1。仅用于面试 Demo，不含真实商业秘密。
用法：python make_corpus.py
"""
import os
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import SimpleDocTemplate, Paragraph

BASE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(BASE, "data")
os.makedirs(DATA, exist_ok=True)

# ---------------- 产品手册.docx ----------------
def build_docx():
    doc = Document()
    doc.add_heading("智能会议平板 X1 产品手册", level=0)
    doc.add_heading("一、产品概述", level=1)
    doc.add_paragraph("云析科技智能会议平板 X1 是面向中大型会议室的一体化协作终端，"
                      "集 4K 触控大屏、白板协作、视频会议与无线传屏于一体，"
                      "适配日常会议、培训与远程协作场景。")
    doc.add_heading("二、硬件规格", level=1)
    doc.add_paragraph("屏幕尺寸：提供 65 / 75 / 86 英寸三种规格，均支持 4K（3840×2160）分辨率、"
                      "20 点触控、防眩光钢化玻璃。")
    doc.add_paragraph("计算单元：标准版 4GB+32GB；Pro 版 8GB+128GB；操作系统为 Android 11，"
                      "可安装第三方协作应用。")
    doc.add_paragraph("接口：HDMI×2（输入/输出各一）、USB-A×3、USB-C（全功能，支持一线连笔记本并反向充电）、"
                      "RJ45 千兆网口、Wi-Fi 6、蓝牙 5.1。")
    doc.add_heading("三、软件功能", level=1)
    doc.add_paragraph("白板协作：支持多人同屏书写、图形识别、扫码带走会议纪要。")
    doc.add_paragraph("视频会议：内置腾讯会议、Zoom、飞书会议、Teams 客户端，支持摄像头与阵列麦。")
    doc.add_paragraph("无线传屏：Windows 安装「云析传屏助手」、macOS 通过 AirPlay 即可一键投屏。")
    doc.add_paragraph("多端同步：会议笔记与白板可同步至手机端与云端空间。")
    doc.add_heading("四、技术对接", level=1)
    doc.add_paragraph("开放平台提供 REST API、Webhook 与 SDK（Python / Java / Node），"
                      "可对接企业微信、飞书、OA 与日历系统，实现会议预约自动拉起、签到数据回写。")
    doc.add_paragraph("支持标准 WebRTC，浏览器内可直接发起会议；支持私有化部署，"
                      "会议媒体与录制数据落在本企业内网，满足数据不出域合规要求。")
    doc.add_heading("五、适用场景", level=1)
    doc.add_paragraph("适用于会议室、培训室、远程协作与门店展示；金融、政企客户可走私有化版本满足合规。")
    doc.add_heading("六、配置与价格", level=1)
    doc.add_paragraph("标准版（65 寸，4+32G）指导价 ¥6,999；Pro 版（86 寸，8+128G，含三年免费上门）"
                      "指导价 ¥19,999。上述为含税价，大客户可走集采折扣。")
    out = os.path.join(DATA, "产品手册.docx")
    doc.save(out)
    print("已生成", out)

# ---------------- 售后政策.pdf ----------------
def build_pdf():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    out = os.path.join(DATA, "售后政策.pdf")
    doc = SimpleDocTemplate(out, pagesize=A4,
                            title="智能会议平板 X1 售后政策")
    ss = getSampleStyleSheet()
    h = ParagraphStyle("h", parent=ss["Heading2"], fontName="STSong-Light")
    p = ParagraphStyle("p", parent=ss["BodyText"], fontName="STSong-Light")
    flow = [
        Paragraph("智能会议平板 X1 售后政策", h),
        Paragraph("一、保修范围", h),
        Paragraph("整机保修 2 年，主要部件（含屏幕）保修 3 年。保修期自发票开具之日起算。"
                  "保修期内非人为故障免费维修或换新。", p),
        Paragraph("二、报修流程", h),
        Paragraph("1）拨打服务热线 400-820-0000 或登录「云析服务」小程序提交工单；"
                  "2）客服远程诊断，可解决则在线指导；3）需上门的，城市区域 24 小时内响应，"
                  "48 小时内上门。", p),
        Paragraph("三、退换货政策", h),
        Paragraph("自签收起 7 天内无理由退货，15 天内性能故障免费换新。退换需保持外观完好、"
                  "配件齐全。人为损坏不在退换范围。", p),
        Paragraph("四、收费标准", h),
        Paragraph("过保后提供有偿维修：上门检测费 ¥100/次（维修则不收），备件按官方价格收取，"
                  "屏幕总成更换为例约 ¥2,800。", p),
        Paragraph("五、服务时效", h),
        Paragraph("全国主要城市提供上门服务，地级市以上 48 小时上门；偏远地区通过寄修处理，"
                  "往返物流由云析承担。", p),
    ]
    doc.build(flow)
    print("已生成", out)

if __name__ == "__main__":
    build_docx()
    build_pdf()
